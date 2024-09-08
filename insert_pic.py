"""
============================
   File Name：     insert_pic
   Description :   
   Author :       师启源
   Date：          2024/9/8
   Notes：
   

============================
"""
import os
import re
from docx import Document
from docx.shared import Inches


class TooManyImagesError(Exception):
    """自定义异常，用于处理文件夹中的图片过多的情况（最多4张图片-身份证2张或1张、毕业证、学位证）"""

    def __init__(self, name, count):
        self.name = name
        self.count = count
        super().__init__(f"当前文件夹 '{name}' 中的图片数量为 {count} 张，超过4张限制，请手动删除非目标图片文件后重新执行。")


class ResumeImageProcessor:
    def __init__(self, root_folder, doc_path):
        self.root_folder = root_folder
        self.doc_path = doc_path
        self.supported_image_extensions = ['.jpg', '.jpeg', '.png']

    def print_document_paragraphs(self, doc):
        print("打印文档段落内容:")
        for para in doc.paragraphs:
            print(para.text)

    def rename_files_in_folder(self):
        for subfolder in os.listdir(self.root_folder):
            subfolder_path = os.path.join(self.root_folder, subfolder)

            if os.path.isdir(subfolder_path):
                print(f"正在处理文件夹: {subfolder}")
                images = [f for f in os.listdir(subfolder_path) if
                          any(f.lower().endswith(ext) for ext in self.supported_image_extensions)]

                # 判断图片数量是否大于 4
                if len(images) > 4:
                    name = re.findall(r'([\u4e00-\u9fa5]+)', subfolder)[0]  # 提取姓名
                    raise TooManyImagesError(name, len(images))

                # 如果图片数量不超过 4，开始重命名操作
                id_card_count = 0
                for filename in images:
                    old_file_path = os.path.join(subfolder_path, filename)

                    if '学位证' in filename or '学位' in filename:
                        new_file_name = "学位证书.jpg"
                        new_file_path = os.path.join(subfolder_path, new_file_name)
                        os.rename(old_file_path, new_file_path)
                        print(f"{filename} 重命名为 {new_file_name}")

                    elif '毕业证' in filename:
                        new_file_name = "毕业证.jpg"
                        new_file_path = os.path.join(subfolder_path, new_file_name)
                        os.rename(old_file_path, new_file_path)
                        print(f"{filename} 重命名为 {new_file_name}")

                    elif '身份证' in filename:
                        id_card_count += 1
                        new_file_name = f"身份证{id_card_count:02d}.jpg"
                        new_file_path = os.path.join(subfolder_path, new_file_name)
                        os.rename(old_file_path, new_file_path)
                        print(f"{filename} 重命名为 {new_file_name}")

    def process_images_and_insert(self, doc):
        for subfolder in os.listdir(self.root_folder):
            subfolder_path = os.path.join(self.root_folder, subfolder)

            if os.path.isdir(subfolder_path):
                print(f"正在处理文件夹: {subfolder}")
                name = re.findall(r'([\u4e00-\u9fa5]+)', subfolder)[0]
                print(f"匹配到的姓名: {name}")

                images = {'身份证01': None, '身份证02': None, '毕业证': None, '学位证书': None}

                for filename in os.listdir(subfolder_path):
                    if any(filename.lower().endswith(ext) for ext in self.supported_image_extensions):
                        file_key = filename.split('.')[0]
                        if file_key in images:
                            images[file_key] = os.path.join(subfolder_path, filename)

                print(f"找到的图片文件: {images}")

                for i, para in enumerate(doc.paragraphs):
                    para_text_cleaned = para.text.strip()

                    if name in para_text_cleaned:
                        print(f"在文档中找到匹配段落: {para_text_cleaned}")
                        self.delete_and_insert_images(doc, i, images)

    def delete_and_insert_images(self, doc, start_index, images):
        # 找到 "身份证：", "毕业证：", "学位证：" 的段落并删除
        delete_start = None
        delete_end = None
        for i in range(start_index + 1, len(doc.paragraphs)):
            if '身份证：' in doc.paragraphs[i].text:
                delete_start = i
            if '学位证：' in doc.paragraphs[i].text:
                delete_end = i
                break

        if delete_start is not None and delete_end is not None:
            for i in range(delete_start, delete_end + 1):
                doc.paragraphs[i]._element.clear()

            self.insert_image(doc.paragraphs[delete_start], images['身份证01'], label="身份证：")
            if images['身份证02']:
                self.insert_image(doc.paragraphs[delete_start + 1], images['身份证02'])

            self.insert_image(doc.paragraphs[delete_start + 2], images['毕业证'], label="毕业证：")
            self.insert_image(doc.paragraphs[delete_start + 3], images['学位证书'], label="学位证：")

    def insert_image(self, paragraph, image_path, label=None):
        if image_path:
            if label:
                paragraph.add_run(label + '\n')
            print(f"插入图片: {image_path}")
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(2))
            paragraph.add_run('\n')

    def save_document(self, doc, output_doc_path):
        doc.save(output_doc_path)
        print(f"文档已保存为 {output_doc_path}")

    def process(self):
        self.rename_files_in_folder()
        doc = Document(self.doc_path)
        self.print_document_paragraphs(doc)
        self.process_images_and_insert(doc)
        self.save_document(doc, '简历插入图片_修改后.docx')


if __name__ == "__main__":
    root_folder = '/Users/shiqiyuan/简历插入图片/简历图片/'  # 图片文件夹路径，windows和mac路径不同，请本地copy后修改
    doc_path = '简历插入图片1.docx'  # Word 文档路径

    try:
        processor = ResumeImageProcessor(root_folder, doc_path)
        processor.process()
    except TooManyImagesError as e:
        print(e)